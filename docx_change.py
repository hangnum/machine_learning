import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# --- 配置参数映射表 ---
FONT_SIZES = {
    '小二': 18,
    '四号': 14,
    '小四': 12,
}

def set_font(run, font_name, size_name, bold=False):
    """
    通用函数：强制设置一段文字的中西文字体、字号和加粗
    """
    run.font.name = 'Times New Roman'  # 西文设置
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 中文强制设置
    run.font.size = Pt(FONT_SIZES[size_name])
    run.font.bold = bold

def set_paragraph_format(paragraph, spacing_val=23, align=None):
    """
    通用函数：设置段落行距、对齐方式，并强制清除段前段后间距
    """
    pf = paragraph.paragraph_format
    
    # 1. 核心修改：强制段前段后间距为 0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    # 2. 设置固定行距 23磅
    if spacing_val:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(spacing_val)
    
    # 3. 设置对齐
    if align is not None:
        paragraph.alignment = align
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐

def auto_format_thesis(input_path, output_path):
    doc = Document(input_path)
    
    # 1. 全局页面设置 (页边距普通: 上下2.54cm, 左右3.18cm)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 2. 遍历所有段落
    paragraphs = doc.paragraphs
    
    # 定义正则匹配模式
    p_h1 = re.compile(r'^一、')       # 一级标题
    p_h2 = re.compile(r'^（[一二三四五六七八九十]+）') # 二级标题
    p_h3 = re.compile(r'^\d+、')      # 三级标题
    
    # 状态标记
    is_abstract_area = False
    is_reference_area = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        
        # 保留原有空行：如果原文档有空行（且不是因为格式造成的视觉空行），需要保留
        # 这里给空行也设置格式，确保空行的高度也是严格的23磅
        if not text:
            set_paragraph_format(p, spacing_val=23) 
            continue 

        # --- A. 识别特殊区域 ---
        
        # 1. 论文题目 (简单逻辑：前5行，字数少，非目录)
        if i < 5 and len(text) < 30 and "题目" not in text and "目录" not in text:
             # 假设这是论文主标题
             p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
             set_paragraph_format(p, spacing_val=23, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
             for run in p.runs:
                 set_font(run, '黑体', '小二', bold=False)
             continue

        # 2. 目录 (单独处理)
        if text == "目录" or text == "目 录":
            set_paragraph_format(p, spacing_val=23, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
            for run in p.runs:
                set_font(run, '黑体', '小二', bold=False)
            continue

        # 3. 摘要和关键词
        if text.startswith("摘") and "要" in text:
            is_abstract_area = True
            p.clear() 
            # 重新构建内容以应用混合字体
            run_title = p.add_run("摘  要：")
            set_font(run_title, '黑体', '小四', bold=True)
            
            content_text = text.split("：")[-1] if "：" in text else text
            run_content = p.add_run(content_text)
            set_font(run_content, '宋体', '小四')
            
            # 这里的对齐方式根据习惯通常是两端对齐
            set_paragraph_format(p, spacing_val=23)
            continue
            
        if text.startswith("关键词"):
            p.clear()
            run_title = p.add_run("关键词：")
            set_font(run_title, '黑体', '小四', bold=True)
            
            content_text = text.split("：")[-1] if "：" in text else text
            run_content = p.add_run(content_text)
            set_font(run_content, '宋体', '小四')
            
            set_paragraph_format(p, spacing_val=23)
            is_abstract_area = False 
            continue

        # 4. 参考文献
        if text == "参考文献" or text == "参考文献：":
            is_reference_area = True
            set_paragraph_format(p, spacing_val=23)
            for run in p.runs:
                set_font(run, '黑体', '小四', bold=True)
            continue
            
        if is_reference_area:
            # 参考文献内容
            set_paragraph_format(p, spacing_val=23) # 保持行距一致
            for run in p.runs:
                set_font(run, '宋体', '小四')
            continue

        # --- B. 识别正文标题 ---

        # 一级标题：黑体4号
        if p_h1.match(text):
            set_paragraph_format(p, spacing_val=23, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            for run in p.runs:
                set_font(run, '黑体', '四号', bold=False)
            continue

        # 二级标题：宋体小4号，加粗
        if p_h2.match(text):
            set_paragraph_format(p, spacing_val=23, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            for run in p.runs:
                # 依然按你正文要求：宋体小四加粗
                set_font(run, '宋体', '小四', bold=True)
            continue

        # 三级标题：宋体小4号
        if p_h3.match(text):
            set_paragraph_format(p, spacing_val=23, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            for run in p.runs:
                set_font(run, '宋体', '小四', bold=False)
            continue

        # --- C. 普通正文 ---
        # 宋体小4号，行距23磅，首行缩进通常需要（这里根据需要可以加首行缩进）
        # 如果需要首行缩进2字符，可以在 set_paragraph_format 里加 pf.first_line_indent = Pt(24)
        set_paragraph_format(p, spacing_val=23)
        for run in p.runs:
            set_font(run, '宋体', '小四')

    doc.save(output_path)
    print(f"排版完成（段前段后间距已清零），已保存为: {output_path}")

# --- 运行部分 ---
auto_format_thesis('20250203010-王文韬-人工智能基础报告（期末）.docx', '20250203010-王文韬-人工智能基础报告（期末）.docx')